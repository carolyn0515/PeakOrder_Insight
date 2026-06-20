from __future__ import annotations

import csv
import math
import statistics
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "report" / "peakorder-insight-aws-cloud-computing-report.docx"
LATENCY_CSV = ROOT / "src" / "outputs" / "aws" / "kinesis_latency_latest.csv"
PUBLISH_CSV = ROOT / "src" / "outputs" / "aws" / "kinesis_publish_peak_replay_latest.csv"
ASSET_DIR = ROOT / "docs" / "report" / "assets"
PEAK_PROFILE_CHART = ASSET_DIR / "peak_order_hourly_profile.png"
ARCHITECTURE_OVERVIEW = ASSET_DIR / "peakorder_architecture_overview.png"
TABLE_COUNTER = 0

SCREENSHOTS = [
    (
        "Figure A-1. Frontend PM KPI dashboard",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-5a9f071c-fb85-455b-a03a-43661a53e1dd.png",
        "PeakOrder Insight의 PM 관점 dashboard이다. 총 주문 수, 피크 윈도우 주문 수, 매출, peak alert, decision freshness가 함께 표시되어 lakehouse 결과가 제품/운영 의사결정 지표로 전환되었음을 보여준다.",
    ),
    (
        "Figure A-2. Kinesis stream monitoring and throughput evidence",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-c3823f65-a4bf-40eb-90cf-b5e88078e2c4.png",
        "Kinesis Data Streams의 monitoring 화면이다. IncomingRecords, Incoming data, PutRecords latency, throughput exceeded/throttling 지표가 표시되어 피크 replay가 실제 AWS streaming layer에 부하로 반영되었음을 증명한다.",
    ),
    (
        "Figure A-3. S3 raw bucket top-level structure",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-fe3c22a8-7bc7-4839-8c86-620d47c1d2f4.png",
        "S3 raw bucket의 최상위 구조이다. inventory, orders, products가 분리되어 있고, 이 중 orders 영역이 주문 이벤트 원천 데이터의 landing zone으로 사용된다.",
    ),
    (
        "Figure A-4. S3 raw order event JSONL landing path",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-a2ac8d6c-82b7-43cc-acc1-397fe35c3d46.png",
        "날짜 파티션 `dt=2026-06-20` 아래에 `order_events.jsonl`이 저장된 화면이다. 원천 주문 이벤트가 S3 raw zone에 보존되어 EMR validation과 재처리의 기준점으로 사용될 수 있음을 보여준다.",
    ),
    (
        "Figure A-5. S3 lakehouse bucket structure",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-d3074ccf-dd44-466f-90d8-43e60d2171b2.png",
        "S3 lakehouse bucket의 구조이다. checkpoints, exports, jars, jobs, paimon 경로가 분리되어 Spark job 자산, Paimon runtime jar, lakehouse table state, serving export가 같은 bucket 안에서 역할별로 관리된다.",
    ),
    (
        "Figure A-6. S3 JSON dashboard exports",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-45c0910c-0a46-4bb1-a738-e09de7fee90c.png",
        "S3 lakehouse의 JSON export 경로이다. frontend dashboard가 빠르게 읽을 수 있도록 product demand, store pressure, peak alerts 등의 serving data가 JSON 형태로 materialize되었음을 보여준다.",
    ),
    (
        "Figure A-7. S3 Parquet analytics exports",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-09327cb3-1dfe-4047-b167-37985a2efdfe.png",
        "S3 lakehouse의 Parquet export 경로이다. 동일한 serving 결과를 columnar format으로 저장해 Glue/Athena-style 분석이나 후속 데이터 처리에 적합하게 만든 상태를 보여준다.",
    ),
    (
        "Figure A-8. Glue database and serving tables",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-2ac4c482-54c1-43a8-b97c-4a564e3ea9ad.png",
        "Glue Data Catalog의 `peakorder_insight_dev` database 화면이다. raw table과 Parquet serving table 5개가 등록되어 S3 파일이 관리 가능한 table metadata로 정리되었음을 보여준다.",
    ),
    (
        "Figure A-9. Glue raw order table schema",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-3e8f3ae5-f0ea-4886-8b34-235ad32576fe.png",
        "Glue의 `raw_order_events_json` table 상세 화면이다. S3 raw orders 경로와 JSON schema가 catalog에 등록되어 raw event를 AWS table 관점에서 검증할 수 있는 상태임을 나타낸다.",
    ),
    (
        "Figure A-10. Glue peak alert serving table schema",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-cb67a254-f048-4d8e-8cde-adb13f0a5305.png",
        "`peak_order_alerts` serving table의 schema 화면이다. store_id, hour_start, order_count, baseline_order_count, pressure_ratio, severity가 포함되어 PM/SRE alert queue의 데이터 근거를 제공한다.",
    ),
    (
        "Figure A-11. CloudWatch log groups",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-ef65eb32-7da4-4f1d-a9b2-7d38ffa1bf7c.png",
        "CloudWatch Logs의 log group 목록이다. `/aws/emr-serverless`와 프로젝트 pipeline log group이 존재해 EMR job 실행 로그와 troubleshooting evidence를 AWS 관찰성 계층에서 확인할 수 있다.",
    ),
    (
        "Figure A-12. CloudWatch peak ingestion alarm",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-82ecacc7-298a-4885-8bad-1a08df580e4e.png",
        "CloudWatch Alarm 화면이다. Kinesis `IncomingRecords >= 40,000 / 1 minute` 기준의 peak ingestion alarm이 설정되어 있고, 피크 구간 datapoint가 threshold를 넘는 것을 확인할 수 있다.",
    ),
    (
        "Figure A-13. VPC resource map",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-d83ca8ed-a0bc-43fb-8b64-47020ecdc845.png",
        "프로젝트 VPC resource map이다. private subnets, route table, S3 gateway endpoint가 연결되어 EMR Serverless processing이 private network 기반으로 S3에 접근하는 구조를 보여준다.",
    ),
    (
        "Figure A-14. VPC endpoints for S3 and CloudWatch Logs",
        "/var/folders/g8/gwqmqng10_g3h7m8r02yq1qc0000gn/T/codex-clipboard-b0be6c26-d8fb-498f-9c0e-5e5d36ace42d.png",
        "VPC endpoint 목록 화면이다. S3 gateway endpoint와 CloudWatch Logs interface endpoint가 모두 Available 상태이며, private subnet의 Spark job이 S3와 Logs에 접근할 수 있는 네트워크 조건을 충족한다.",
    ),
]


def pct(vals: list[float], p: float) -> float:
    if not vals:
        return 0.0
    vals = sorted(vals)
    k = (len(vals) - 1) * p / 100
    lo = math.floor(k)
    hi = math.ceil(k)
    if lo == hi:
        return vals[int(k)]
    return vals[lo] * (hi - k) + vals[hi] * (k - lo)


def load_latency_stats() -> dict[str, dict[str, float]]:
    rows: list[dict[str, object]] = []
    if not LATENCY_CSV.exists():
        return {}
    with LATENCY_CSV.open(encoding="utf-8") as f:
        for r in csv.DictReader(f):
            try:
                r["pub_ms"] = float(r["publish_to_read_latency_ms"])
                r["arr_ms"] = float(r["kinesis_arrival_latency_ms"])
            except (TypeError, ValueError):
                continue
            rows.append(r)

    peak_hours = {"12:00", "13:00", "18:00", "19:00"}
    shoulder_hours = {"08:00", "09:00", "10:00", "11:00", "14:00", "15:00", "16:00", "17:00", "20:00", "21:00"}
    off_hours = {"00:00", "01:00", "02:00", "03:00", "04:00", "05:00", "06:00", "07:00", "22:00", "23:00"}

    def summarize(label: str, subset: list[dict[str, object]]) -> dict[str, float]:
        vals = [float(r["pub_ms"]) for r in subset]
        arr = [float(r["arr_ms"]) for r in subset]
        return {
            "label": label,
            "records": len(vals),
            "mean": statistics.mean(vals) if vals else 0.0,
            "median": statistics.median(vals) if vals else 0.0,
            "std": statistics.pstdev(vals) if len(vals) > 1 else 0.0,
            "p95": pct(vals, 95),
            "p99": pct(vals, 99),
            "max": max(vals) if vals else 0.0,
            "arrival_mean": statistics.mean(arr) if arr else 0.0,
            "arrival_p95": pct(arr, 95),
        }

    return {
        "all": summarize("All records", rows),
        "peak": summarize("Peak hours", [r for r in rows if r["event_hour"] in peak_hours]),
        "shoulder": summarize("Shoulder hours", [r for r in rows if r["event_hour"] in shoulder_hours]),
        "off": summarize("Off-peak hours", [r for r in rows if r["event_hour"] in off_hours]),
    }


def load_publish_stats() -> dict[str, object]:
    rows: list[dict[str, object]] = []
    if not PUBLISH_CSV.exists():
        return {}
    with PUBLISH_CSV.open(encoding="utf-8") as f:
        for r in csv.DictReader(f):
            r["elapsed_seconds"] = float(r["elapsed_seconds"])
            r["published_records"] = int(r["published_records"])
            r["failed_records"] = int(r["failed_records"])
            r["records_per_second"] = float(r["records_per_second"])
            rows.append(r)
    if not rows:
        return {}
    prev_elapsed = 0.0
    prev_records = 0
    by_hour: dict[str, dict[str, float]] = {}
    for r in rows:
        h = str(r["event_hour"])
        delta_records = int(r["published_records"]) - prev_records
        delta_seconds = max(float(r["elapsed_seconds"]) - prev_elapsed, 0.0001)
        d = by_hour.setdefault(h, {"records": 0, "seconds": 0.0, "max_rps": 0.0})
        d["records"] += delta_records
        d["seconds"] += delta_seconds
        d["max_rps"] = max(d["max_rps"], delta_records / delta_seconds)
        prev_elapsed = float(r["elapsed_seconds"])
        prev_records = int(r["published_records"])
    return {
        "total": rows[-1]["published_records"],
        "elapsed": rows[-1]["elapsed_seconds"],
        "avg_rps": rows[-1]["records_per_second"],
        "failed_sum": sum(int(r["failed_records"]) for r in rows),
        "batches": len(rows),
        "by_hour": by_hour,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    caption: str | None = None,
) -> None:
    global TABLE_COUNTER
    TABLE_COUNTER += 1
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="0B2545")
        set_cell_shading(hdr[i], "F2F4F7")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Inches(widths[i])
    caption_text = caption or f"{headers[0]} summary"
    if not caption_text.startswith("Table "):
        caption_text = f"Table {TABLE_COUNTER}. {caption_text}"
    add_caption(doc, caption_text)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    for line in text.strip("\n").splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("1F2937")


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph()


def add_label_line(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string("0B2545")
    label_run.font.size = Pt(10)
    body_run = p.add_run(text)
    body_run.font.size = Pt(10)


def add_decision_block(
    doc: Document,
    technology: str,
    compared_with: str,
    differentiator: str,
    project_fit: str,
    interview_answer: str,
) -> None:
    doc.add_heading(technology, level=2)
    add_label_line(doc, "대안 비교: ", compared_with)
    add_label_line(doc, "기술적 차별점: ", differentiator)
    add_label_line(doc, "프로젝트 요구사항: ", project_fit)
    add_label_line(doc, "면접 답변식 선택 근거: ", interview_answer)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_figure_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("344054")


def add_fit_picture(doc: Document, path: Path, max_width: float = 6.2, max_height: float = 6.9) -> None:
    with Image.open(path) as img:
        px_w, px_h = img.size
    ratio = px_w / px_h if px_h else 1
    width = min(max_width, max_height * ratio)
    height = width / ratio if ratio else max_height
    if height > max_height:
        height = max_height
        width = height * ratio
    doc.add_picture(str(path), width=Inches(width), height=Inches(height))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def create_peak_profile_chart() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    hours = list(range(24))
    values = [1666] * 8 + [5000, 5000, 5000, 10000, 50000, 50000, 5000, 5000, 5000, 10000, 50000, 50000, 10000, 5000, 1666, 1666]
    peak_hours = {12, 13, 18, 19}
    w, h = 1800, 780
    margin_l, margin_r, margin_t, margin_b = 120, 80, 120, 120
    plot_w = w - margin_l - margin_r
    plot_h = h - margin_t - margin_b
    max_y = 55000

    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font = font(44, bold=True)
    sub_font = font(24)
    axis_font = font(22)
    label_font = font(20)
    small_font = font(18)

    d.text((margin_l, 36), "PeakOrder hourly order profile", fill="#0B2545", font=title_font)
    d.text(
        (margin_l, 88),
        "Synthetic order events are intentionally concentrated at lunch and dinner windows to test burst ingestion.",
        fill="#52616F",
        font=sub_font,
    )

    # Background and grid
    d.rounded_rectangle((margin_l, margin_t, w - margin_r, h - margin_b), radius=18, fill="#F6F9FC", outline="#D9E2EC", width=2)
    for yv in [0, 10000, 20000, 30000, 40000, 50000]:
        y = margin_t + plot_h - (yv / max_y) * plot_h
        d.line((margin_l, y, w - margin_r, y), fill="#E1E8F0", width=2)
        d.text((35, y - 13), f"{yv//1000}k", fill="#6B7C8F", font=small_font)

    bar_gap = 10
    bar_w = (plot_w - bar_gap * 23) / 24
    for i, (hour, value) in enumerate(zip(hours, values)):
        x0 = margin_l + i * (bar_w + bar_gap)
        x1 = x0 + bar_w
        y1 = margin_t + plot_h
        y0 = y1 - (value / max_y) * plot_h
        color = "#C93A0A" if hour in peak_hours else "#0F766E"
        d.rounded_rectangle((x0, y0, x1, y1), radius=8, fill=color)
        d.text((x0 + bar_w / 2 - 13, h - margin_b + 28), f"{hour:02d}", fill="#52616F", font=axis_font)

    # Highlight peak windows
    def annotate(hour_start: int, text: str) -> None:
        idx = hour_start
        x = margin_l + idx * (bar_w + bar_gap) + bar_w / 2
        y = margin_t + plot_h - (50000 / max_y) * plot_h - 48
        d.rounded_rectangle((x - 105, y - 42, x + 105, y - 4), radius=10, fill="#FFF3E8", outline="#F97316", width=2)
        d.text((x - 88, y - 36), text, fill="#9A3412", font=label_font)
        d.line((x, y - 4, x, y + 35), fill="#F97316", width=3)

    annotate(12, "Lunch peak")
    annotate(18, "Dinner peak")

    # Legend
    lx, ly = w - 520, 48
    d.rounded_rectangle((lx, ly, lx + 430, ly + 44), radius=10, fill="#FFFFFF", outline="#D9E2EC")
    d.rounded_rectangle((lx + 18, ly + 14, lx + 46, ly + 30), radius=5, fill="#0F766E")
    d.text((lx + 58, ly + 10), "Normal / shoulder", fill="#344054", font=small_font)
    d.rounded_rectangle((lx + 242, ly + 14, lx + 270, ly + 30), radius=5, fill="#C93A0A")
    d.text((lx + 282, ly + 10), "Peak window", fill="#344054", font=small_font)

    d.text((margin_l, h - 46), "Hour of day", fill="#52616F", font=axis_font)
    d.text((12, margin_t + 10), "Orders", fill="#52616F", font=axis_font)

    img.save(PEAK_PROFILE_CHART)
    return PEAK_PROFILE_CHART


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        st = styles[list_style]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "PeakOrder Insight AWS Cloud Computing Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string("555555")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("PeakOrder Insight | AWS Cloud Computing Project").font.size = Pt(9)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PeakOrder Insight")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AWS 기반 피크 주문 실시간 스트리밍 및 Lakehouse 분석\n아키텍처 보고서")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AWS Cloud Computing Course Project")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()
    add_callout(
        doc,
        "Executive Positioning",
        "본 보고서는 단순 웹사이트 제작 결과물이 아니라, 피크 시간대 주문 폭증이라는 운영 문제를 AWS cloud service, Terraform IaC, private networking, IAM permission, streaming observability, serverless Spark processing, lakehouse table format, dashboard serving까지 연결해 검증한 cloud computing 설계 및 구현 보고서이다.",
        fill="E8EEF5",
    )

    meta = [
        ["Project", "PeakOrder Insight"],
        ["Region", "ap-northeast-2 / Asia Pacific Seoul"],
        ["Primary AWS Services", "Kinesis, S3, EMR Serverless, Glue, CloudWatch, VPC, IAM, SNS"],
        ["Infrastructure Method", "Terraform module-based Infrastructure as Code"],
        ["Experiment Scale", "281,660 order events / 200,000 peak-window events"],
        ["Scale Interpretation", "Time-compressed replay baseline, not a production full-scale guarantee"],
        ["Validation Result", "Current replay p95 latency 2.03s overall / 2.06s peak"],
    ]
    add_table(doc, ["Item", "Value"], meta, [1.8, 4.5])
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("목차", level=1)
    entries = [
        "1. 프로젝트 개요와 문제 정의",
        "2. AWS Cloud Computing 관점의 설계 목표",
        "3. 핵심 AWS 기술 선택 근거",
        "4. 전체 아키텍처 상세 설명",
        "5. Terraform 기반 Infrastructure as Code",
        "6. Kinesis 스트리밍 수집과 피크 부하 검증",
        "7. S3 Raw Zone, Lakehouse Zone, Paimon Table Format",
        "8. EMR Serverless Spark 처리 계층",
        "9. Glue Data Catalog와 Serving Table 설계",
        "10. VPC, Private Subnet, Endpoint, IAM 권한 설계",
        "11. CloudWatch 관찰성과 알람 설계",
        "12. Frontend Dashboard와 PM/SRE 의사결정 지표",
        "13. Latency 분석 결과와 구조 타당성 평가",
        "14. 트러블슈팅과 개선 과정",
        "15. 한계, 비용, 운영 개선안",
        "16. 결론",
        "Appendix A. AWS Console Evidence",
    ]
    add_numbered(doc, entries)
    doc.add_page_break()


def build_report() -> None:
    global TABLE_COUNTER
    TABLE_COUNTER = 0
    chart_path = create_peak_profile_chart()
    latency = load_latency_stats()
    publish = load_publish_stats()
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_contents(doc)

    doc.add_heading("1. 프로젝트 개요와 문제 정의", level=1)
    doc.add_paragraph(
        "PeakOrder Insight는 점심과 저녁처럼 주문이 특정 시간대에 몰리는 배달/커머스 서비스를 가정하고, 피크 시간대 주문 폭증을 실시간으로 감지하고 lakehouse에 재처리 가능한 형태로 축적한 뒤, 운영자와 PM이 바로 사용할 수 있는 지표로 제공하는 AWS 기반 데이터 엔지니어링 프로젝트이다."
    )
    doc.add_paragraph(
        "이 프로젝트의 중요한 출발점은 'AWS 서비스를 많이 사용했다'가 아니라 '왜 이 문제에 AWS cloud architecture가 필요한가'이다. 피크 주문 시스템에서는 하루 총 주문량보다 순간적으로 몰리는 주문 밀도가 더 큰 리스크가 된다. 같은 28만 건의 주문이라도 24시간에 균등하게 들어오는 경우와 12시, 13시, 18시, 19시에 집중되는 경우는 capacity planning, store staffing, inventory replenishment, delivery ETA, customer support 부담이 완전히 달라진다."
    )
    add_callout(
        doc,
        "Core thesis",
        "PeakOrder Insight의 핵심은 피크 시간대 주문 폭증을 Kinesis로 빠르게 관찰하고, S3/Paimon/Glue 기반 lakehouse로 재현 가능한 분석 상태를 만들며, CloudWatch와 dashboard를 통해 운영 판단까지 연결하는 것이다.",
    )
    doc.add_heading("1.1 Peak-shaped order event generator 설계", level=2)
    doc.add_paragraph(
        "이 프로젝트의 실험은 단순히 281,660건의 JSON 데이터를 만든 뒤 한 번에 업로드한 것이 아니다. 실제 운영에서 문제가 되는 것은 하루 총 주문량보다 '특정 시간대에 주문이 얼마나 조밀하게 몰리는가'이므로, generator 단계에서부터 주문 이벤트의 시간 분포를 의도적으로 비균등하게 설계했다. 기본 정상 시간대는 시간당 5,000건, 피크 시간대는 peak multiplier 10을 적용해 시간당 50,000건으로 설정했고, 점심 피크 12-13시와 저녁 피크 18-19시에 전체 주문의 대부분이 몰리도록 만들었다."
    )
    doc.add_paragraph(
        "생성된 event는 event_id, order_id, customer_id, store_id, event_type, event_time, items를 포함한다. event_time은 업무일 기준 00시부터 23시까지 분산되지만, volume function은 시간대별로 다르게 적용된다. 11시, 17시, 20시는 피크 전후 ramp 구간으로 정상 시간대의 2배를 넣고, 00-07시와 22-23시는 정상 시간대의 1/3 수준만 생성한다. 이 방식은 실제 매장의 조용한 시간대, 피크 직전 준비 구간, 점심/저녁 피크, 피크 이후 완화 구간을 하나의 synthetic business day 안에 재현하기 위한 설계다."
    )
    doc.add_page_break()
    add_table(
        doc,
        ["Traffic segment", "Event hours", "Generated orders", "Business meaning", "30s compressed rate"],
        [
            ["Off-peak", "00-07, 22-23", "1,666/hour", "새벽/야간 baseline", "약 56 records/sec"],
            ["Shoulder", "08-10, 14-16, 21", "5,000/hour", "일반 영업 시간대", "약 167 records/sec"],
            ["Ramp", "11, 17, 20", "10,000/hour", "피크 직전/직후 상승 구간", "약 333 records/sec"],
            ["Peak", "12, 13, 18, 19", "50,000/hour", "점심/저녁 주문 폭증", "약 1,667 records/sec"],
        ],
        [1.05, 1.25, 1.35, 1.6, 1.35],
        caption="Peak-shaped generator의 시간대별 주문량 설계와 압축 replay 목표 속도",
    )
    doc.add_paragraph(
        "여기서 중요한 실험 설계 포인트는 시간 압축이다. AWS Kinesis, S3, EMR Serverless 같은 managed service는 실제 production 설정에서는 훨씬 큰 처리량까지 확장될 수 있으므로, 수업 프로젝트 환경에서 수천만-수억 건을 장시간 발생시키는 방식은 비용과 권한 측면에서 적절하지 않다. 대신 본 프로젝트는 event-time 기준 1시간을 wall-clock 30초에 replay하도록 설정했다. 즉, 하루치 업무 시간을 실제 24시간이 아니라 약 12분 수준으로 압축해 흘려보낸다. 이 압축률은 3,600초를 30초로 줄인 것이므로 120배이며, 피크 시간대 50,000건은 이론적으로 약 1,667 records/sec의 순간 유입 압력으로 전환된다."
    )
    add_callout(
        doc,
        "Experiment legitimacy",
        "따라서 이 실험은 '281,660건이면 대규모 production 부하다'라고 주장하는 실험이 아니다. 더 정확히는 제한된 데이터셋을 120배 빠른 시간 압축으로 replay하여 피크형 유입 압력, throttling 여부, retry/backoff 동작, p95/p99 latency, CloudWatch alarm 전이를 관찰하는 scaled benchmark이다.",
        fill="FFF6E8",
    )
    add_table(
        doc,
        ["운영 질문", "필요한 데이터", "사용 계층"],
        [
            ["지금 주문이 평소 대비 얼마나 몰리는가?", "시간대별 order count, store baseline", "PM/SRE dashboard"],
            ["어느 매장이 가장 큰 압력을 받는가?", "store-hour pressure ratio", "Peak alert table"],
            ["어떤 상품이 피크 수요를 주도하는가?", "product demand, gross sales", "PM inventory decision"],
            ["실시간 구조가 실제로 빠른가?", "publish-to-read latency", "Architecture validation"],
            ["AWS에서 운영 관찰이 가능한가?", "CloudWatch metrics/logs/alarm", "Observability layer"],
        ],
        [2.2, 2.2, 1.8],
    )

    doc.add_heading("2. AWS Cloud Computing 관점의 설계 목표", level=1)
    doc.add_paragraph(
        "AWS Cloud Computing 수업 프로젝트로서 이 설계는 단순한 데이터 처리 코드보다 클라우드 서비스의 역할 분담, 네트워크 보안, 권한 제어, 운영 관찰성, 확장성을 명확히 보여주는 것이 중요하다. 따라서 각 AWS 서비스는 기능을 나열하기 위해 선택한 것이 아니라, 피크 주문이라는 문제의 특정 요구사항을 해결하기 위해 배치했다."
    )
    add_table(
        doc,
        ["설계 목표", "AWS 구현", "평가 기준"],
        [
            ["Burst ingestion", "Amazon Kinesis Data Streams", "피크 시간대 IncomingRecords 증가와 PutRecords 성공률"],
            ["Durable raw storage", "Amazon S3 raw bucket", "JSONL 원천 이벤트 보존, backfill/replay 가능성"],
            ["Serverless processing", "EMR Serverless Spark", "클러스터 직접 운영 없이 Spark validation/transform 실행"],
            ["Lakehouse state", "Apache Paimon on S3", "최신 상태 table과 분석 table을 S3 위에 유지"],
            ["Metadata management", "AWS Glue Data Catalog", "Raw/serving table schema와 S3 location 관리"],
            ["Observability", "CloudWatch Logs, Metrics, Alarm", "피크 유입, job log, 장애 원인 확인 가능"],
            ["Private networking", "VPC, private subnets, endpoints", "S3/Logs 접근을 private path로 구성"],
            ["Infrastructure reproducibility", "Terraform modules", "동일 리소스 재생성, 변경 추적, destroy 가능"],
        ],
        [1.7, 2.0, 2.6],
    )
    doc.add_paragraph(
        "특히 이 프로젝트는 load balancing을 전통적인 ALB 기반 웹 트래픽 분산으로 구현하지 않았다. 대신 데이터 ingestion 관점의 load balancing, 즉 Kinesis partition key 기반 shard 분산, producer batch/retry, throughput metric 관찰을 중심으로 피크 부하를 다룬다. 웹 서버 앞단의 ALB가 HTTP request를 분산한다면, Kinesis는 order event stream을 shard 단위로 분산 수용하는 cloud-native ingestion buffer 역할을 한다."
    )

    doc.add_heading("3. 핵심 AWS 기술 선택 근거", level=1)
    doc.add_paragraph(
        "교수자가 데이터 엔지니어링 세부 기술에 익숙하지 않더라도 전체 설계를 이해할 수 있도록, 이 섹션에서는 프로젝트에 사용한 핵심 기술을 '무엇인지', '장단점은 무엇인지', '대안이 있었는데 왜 이것을 선택했는지' 관점에서 정리한다. 핵심은 특정 서비스를 많이 쓴 것이 아니라, 피크 주문 문제의 요구사항에 맞게 서비스를 역할별로 배치했다는 점이다."
    )
    add_decision_block(
        doc,
        "Amazon Kinesis Data Streams",
        "SQS, EventBridge, Amazon MSK/Kafka와 비교했다. SQS는 비동기 작업 queue에 가깝고, EventBridge는 event routing에 강하며, MSK/Kafka는 강력하지만 broker 운영 부담이 크다.",
        "Kinesis는 AWS managed streaming service로 shard 기반 ordered stream, replay, consumer iterator, throughput metric을 제공한다. queue처럼 메시지를 단순 소비하는 구조보다 stream position과 peak throughput을 관찰하기 쉽다.",
        "이 프로젝트는 점심/저녁 피크에 주문 이벤트가 몰리는 상황을 재현하고, records/sec, PutRecords latency, throttling, consumer latency를 함께 봐야 한다.",
        "피크 주문 이벤트는 단순 background job이 아니라 시간 순서와 replay가 중요한 stream workload이다. Kafka급 운영 복잡도는 과하고, AWS-native metric과 managed 운영이 중요한 실험이므로 Kinesis가 가장 합리적이다.",
    )
    add_decision_block(
        doc,
        "Amazon S3",
        "RDS, DynamoDB, EFS, local disk와 비교했다. RDS/DynamoDB는 serving database 성격이 강하고 query pattern에 묶인다. EFS/local disk는 데이터 레이크의 장기 내구성, 비용 효율, 서비스 간 공유성이 약하다.",
        "S3는 object storage 기반의 data lake 표준이다. raw JSONL, Parquet export, Paimon warehouse, job assets를 저렴하고 내구성 높은 저장소에 분리해 둘 수 있다.",
        "원천 주문 이벤트를 immutable하게 보관하고, validation logic이나 aggregate 기준이 바뀌면 Spark/Glue/dashboard가 같은 raw data를 다시 처리할 수 있어야 한다.",
        "주문 이벤트를 바로 serving DB에 넣으면 재처리와 감사가 어려워진다. S3 raw zone은 원천 이벤트의 기준점이 되고, lakehouse/export zone은 가공 결과를 분리하므로 데이터 엔지니어링 관점에서 더 적합하다.",
    )
    add_decision_block(
        doc,
        "EMR Serverless",
        "Glue ETL, Lambda, ECS batch worker, EMR on EC2와 비교했다. Lambda는 대용량 Spark transform과 custom lakehouse jar 실행에 부적합하고, EMR on EC2는 제어권은 크지만 cluster 운영 부담이 크다.",
        "EMR Serverless는 Spark job 실행과 AWS managed 운영의 균형을 제공한다. 상시 cluster를 관리하지 않으면서도 Spark, S3, Glue, VPC, IAM, runtime jar 설정을 결합할 수 있다.",
        "JSONL raw events를 검증하고, store-hour pressure, product demand, peak alert, Parquet/JSON export처럼 batch aggregation을 실행해야 한다.",
        "이 프로젝트는 상시 실행 서비스가 아니라 job 단위 lakehouse processing이다. 따라서 EC2 cluster를 계속 운영하기보다 serverless Spark로 실행 비용과 운영 부담을 낮추는 것이 맞다.",
    )
    add_decision_block(
        doc,
        "Apache Paimon",
        "Plain Parquet, Apache Iceberg, Hudi, Delta Lake와 비교했다. Plain Parquet은 단순 파일이라 snapshot/table evolution/current-state 관리가 약하다. Iceberg는 analytic table 표준성이 강하고, Hudi는 upsert/CDC use case에 강하다.",
        "Paimon은 streaming-batch 통합과 최신 상태 table을 lakehouse에 유지하는 관점이 강하다. S3 위에서 manifest, snapshot, data file을 관리해 Spark가 table처럼 읽고 쓸 수 있게 한다.",
        "피크 주문의 최신 store pressure, alert state, dashboard serving snapshot처럼 운영성 current view를 raw event와 분리해 관리해야 한다.",
        "이 프로젝트의 핵심은 raw archive가 아니라 '피크 상황의 최신 운영 상태'를 빠르게 반영하는 것이다. 단순 Parquet dump보다 table state 계층이 필요하므로 Paimon을 state layer로 선택했다. 다만 production 표준성과 생태계만 보면 Iceberg도 강한 대안이라는 점은 함께 인정한다.",
    )
    add_decision_block(
        doc,
        "AWS Glue Data Catalog",
        "Manual S3 path management, Hive Metastore, RDS metadata table과 비교했다. 수동 path 관리는 schema discoverability가 약하고, 별도 metastore/RDS는 운영 대상이 늘어난다.",
        "Glue는 AWS-native metadata catalog로 S3 location, schema, classification을 관리한다. 데이터 자체를 저장하지는 않지만 EMR, Athena, Glue 계열 서비스가 같은 table metadata를 공유하게 한다.",
        "raw JSONL과 serving Parquet export를 table로 등록해 데이터 위치와 schema를 설명 가능하게 만들어야 한다.",
        "S3에 파일만 있으면 데이터셋의 의미가 코드와 문서에 흩어진다. Glue를 사용하면 raw table과 serving table의 schema/location이 catalog에 남아 분석 서비스로 확장하기 쉽다.",
    )
    add_decision_block(
        doc,
        "CloudWatch",
        "Prometheus/Grafana, OpenSearch, custom log files와 비교했다. Grafana stack은 세밀한 observability에 강하지만 별도 운영 구성이 필요하고, OpenSearch는 log search에는 강하지만 metric alarm의 기본 통합은 CloudWatch가 단순하다.",
        "CloudWatch는 AWS managed service의 기본 metric/log/alarm plane이다. Kinesis, EMR Serverless, Logs, Alarm을 별도 agent 없이 AWS 계정 안에서 연결할 수 있다.",
        "Kinesis throughput, PutRecords latency, EMR logs, peak ingestion alarm을 AWS-native 방식으로 관찰해야 한다.",
        "관찰 대상이 대부분 AWS managed service이므로 1차 observability plane은 CloudWatch가 자연스럽다. production에서 더 정교한 tracing이 필요하면 Grafana/OpenSearch를 붙이되, 현재 프로젝트에서는 비용과 운영 복잡도 대비 CloudWatch가 충분하다.",
    )
    add_decision_block(
        doc,
        "VPC Endpoint",
        "Public subnet, NAT Gateway, Internet Gateway와 비교했다. Public subnet은 쉽지만 compute가 public routing에 노출되고, NAT Gateway는 private outbound를 제공하지만 비용이 든다.",
        "VPC endpoint는 private subnet의 compute가 AWS service에 private path로 접근하게 한다. S3는 gateway endpoint로, CloudWatch Logs는 interface endpoint로 연결할 수 있다.",
        "EMR Serverless processing이 S3 raw/lakehouse와 CloudWatch Logs에 접근해야 하지만, 처리 계층은 private network 기반으로 설명되어야 한다.",
        "S3 gateway endpoint는 대량 S3 접근을 private route로 처리하고 NAT 비용을 줄일 수 있다. Logs interface endpoint는 private subnet에서 job log delivery를 가능하게 하므로 private processing과 AWS service access를 동시에 만족한다.",
    )
    add_decision_block(
        doc,
        "IAM Role / Policy",
        "Root/static access key, broad AdministratorAccess, service-linked default에만 의존하는 방식과 비교했다. 정적 키와 root 권한은 유출 위험이 크고, AdministratorAccess는 최소 권한 원칙에 맞지 않는다.",
        "IAM role은 AWS service가 필요한 리소스에 임시 권한으로 접근하게 한다. S3, Glue, Logs, EMR API 권한을 job role 단위로 제한할 수 있다.",
        "EMR job이 S3, Glue, Logs에 접근해야 하지만 권한 범위는 프로젝트 리소스로 제한되어야 한다.",
        "데이터 플랫폼은 코드보다 권한 경계에서 자주 실패한다. IAM role/policy를 명시적으로 설계하면 job 실행 권한, catalog 접근, log 기록 권한을 분리해 troubleshooting과 보안 설명이 가능하다.",
    )
    add_decision_block(
        doc,
        "Terraform",
        "Manual console setup, CloudFormation, AWS CDK와 비교했다. 콘솔 수동 생성은 재현성과 변경 추적이 약하고, CloudFormation은 AWS-native지만 문법이 장황하며, CDK는 프로그래밍 추상화가 강하다.",
        "Terraform은 provider 기반 IaC로 module화, plan, apply, destroy workflow가 직관적이다. 여러 AWS 리소스를 동일한 naming convention과 environment 구조로 관리하기 좋다.",
        "VPC, S3, Kinesis, Glue, EMR, IAM, CloudWatch 리소스를 반복 생성/검증/삭제할 수 있어야 한다.",
        "이 프로젝트는 한 번 콘솔에서 만든 리소스를 보여주는 것이 아니라 architecture를 코드로 재현해야 한다. Terraform은 변경 전 plan 검토와 실습 후 destroy가 명확해 비용 관리와 포트폴리오 재현성 측면에서 적합하다.",
    )
    doc.add_paragraph(
        "기술 면접에서 이 선택을 설명한다면 핵심은 '유명한 서비스를 골랐다'가 아니라 workload 특성과 운영 제약의 매칭이다. 주문 이벤트는 비동기 queue가 아니라 시간 순서와 replay가 중요한 stream이므로 Kinesis를 선택했고, 원천 이벤트는 transactional database보다 S3 raw zone에 보관하는 것이 재처리와 감사에 유리하다. Spark 처리는 상시 cluster가 아니라 job 단위 처리이므로 EMR Serverless가 맞고, 최신 운영 상태는 단순 Parquet 파일보다 lakehouse table format으로 관리하는 편이 자연스럽다. 마지막으로 Glue, IAM, VPC endpoint, CloudWatch, Terraform은 각각 metadata, permission boundary, private routing, observability, reproducibility라는 cloud architecture의 필수 운영면을 담당한다."
    )
    add_callout(
        doc,
        "Teaching-oriented interpretation",
        "이 프로젝트의 기술 선택은 '데이터 엔지니어링에서 유명한 도구를 썼다'가 아니라, AWS cloud에서 부하 수집, 저장, 처리, 권한, 네트워크, 관찰성, 비용 정리까지 하나의 시스템으로 설계했다는 점을 보여준다.",
        fill="E8EEF5",
    )

    doc.add_heading("4. 전체 아키텍처 상세 설명", level=1)
    doc.add_paragraph("아키텍처는 크게 여섯 개 레이어로 이해할 수 있다.")
    if ARCHITECTURE_OVERVIEW.exists():
        add_fit_picture(doc, ARCHITECTURE_OVERVIEW, max_width=6.4, max_height=2.8)
        add_caption(doc, "Figure 1. PeakOrder Insight end-to-end AWS architecture")
        doc.add_paragraph(
            "위 다이어그램은 PeakOrder Insight의 전체 데이터 흐름을 한눈에 보여준다. Local/Simulation Layer의 peak-shaped order event generator는 동일한 주문 이벤트를 두 갈래로 사용한다. 첫 번째 경로는 S3 Raw Bucket에 JSONL 원천 데이터를 저장하는 raw data lake 경로이고, 두 번째 경로는 Kinesis Data Streams로 이벤트를 replay하여 streaming ingestion과 latency 관찰을 수행하는 경로이다. EMR Serverless Spark는 private subnet 안에서 raw 데이터를 읽어 validation, Paimon table bootstrap/load, pressure detection, serving export materialization을 실행한다. 처리 결과는 S3 lakehouse bucket의 Paimon warehouse와 JSON/Parquet serving exports로 저장되고, Glue Data Catalog는 raw 및 serving table의 schema와 S3 location metadata를 관리한다. CloudWatch는 Kinesis stream metrics, EMR logs, alarm 상태를 통해 운영 관찰성을 제공하며, dashboard layer는 serving exports와 catalog metadata를 바탕으로 PM/SRE/Leadership 관점의 지표를 표시한다."
        )
    add_table(
        doc,
        ["Layer", "구성 요소", "역할"],
        [
            ["Simulation/Event Source", "peak-shaped producer", "피크 주문 이벤트를 생성하고 Kinesis로 재생"],
            ["Streaming Ingestion", "Kinesis Data Streams", "실시간 수집, throughput metric, consumer latency 검증"],
            ["Raw Data Lake", "S3 raw bucket", "원천 JSONL 보관, audit/replay/backfill 기준점"],
            ["Processing", "EMR Serverless Spark", "검증, 변환, Paimon 적재, alert 생성, export 생성"],
            ["Lakehouse/Catalog", "S3 lakehouse, Paimon, Glue", "table state, serving export, schema metadata 관리"],
            ["Observability/Serving", "CloudWatch, dashboard", "운영 metric/log/alarm과 역할별 의사결정 화면 제공"],
        ],
        [1.5, 2.1, 2.7],
    )
    doc.add_paragraph(
        "데이터 흐름은 왼쪽에서 오른쪽으로 읽으면 된다. Producer가 peak-shaped order events를 생성해 Kinesis에 PutRecords로 전송한다. 동시에 같은 이벤트 데이터셋은 S3 raw zone에 JSONL 형태로 저장되어 EMR Serverless의 입력이 된다. EMR Serverless는 private VPC 환경에서 raw JSONL을 읽고 Spark job을 실행해 데이터 품질을 검증하고, Paimon lakehouse table과 JSON/Parquet serving export를 생성한다. Glue Data Catalog는 이 결과의 schema와 S3 location을 관리한다. CloudWatch는 Kinesis metrics, peak ingestion alarm, EMR logs를 통해 운영 관찰성을 제공한다."
    )

    doc.add_heading("5. Terraform 기반 Infrastructure as Code", level=1)
    doc.add_paragraph(
        "이 프로젝트는 AWS 콘솔에서 수동으로 리소스를 만든 것이 아니라 Terraform module 구조를 통해 재현 가능한 인프라를 구성했다. Cloud Computing 수업 관점에서 IaC는 중요한 평가 지점이다. Terraform을 사용하면 리소스 생성 과정이 코드로 남고, 어떤 IAM policy가 부여되었는지, 어떤 VPC endpoint가 생성되었는지, 어떤 S3 bucket과 Glue table이 연결되는지 추적할 수 있다."
    )
    add_table(
        doc,
        ["Terraform Module", "담당 리소스", "설계 의미"],
        [
            ["network", "VPC, private subnets, route tables, S3 endpoint, Logs endpoint", "private processing과 AWS service access의 기반"],
            ["storage", "S3 raw/lakehouse buckets", "raw zone과 lakehouse zone 분리"],
            ["streaming", "Kinesis stream", "피크 event ingestion layer"],
            ["glue", "Glue database and external tables", "metadata catalog 관리"],
            ["emr", "EMR Serverless application", "serverless Spark processing runtime"],
            ["iam", "pipeline role and policies", "least-privilege 기반 접근 제어"],
            ["observability", "CloudWatch log groups, alarm, SNS topic", "운영 관찰성과 peak alerting"],
        ],
        [1.45, 2.4, 2.45],
    )
    doc.add_paragraph(
        "Terraform을 사용한 장점은 세 가지이다. 첫째, 인프라가 문서가 아니라 실행 가능한 코드로 남는다. 둘째, 수업이나 포트폴리오 제출 이후에도 같은 구조를 다른 AWS 계정에 재현할 수 있다. 셋째, 실습이 끝난 뒤 destroy를 통해 비용이 발생하는 리소스를 정리할 수 있다. 특히 Kinesis shard, EMR Serverless application, S3 bucket, CloudWatch log retention은 장기 방치 시 비용이 누적될 수 있으므로 IaC 기반 정리가 중요하다."
    )

    doc.add_heading("6. Kinesis 스트리밍 수집과 피크 부하 검증", level=1)
    doc.add_paragraph(
        "Kinesis Data Streams는 이 프로젝트의 실시간 ingestion layer이다. 주문 이벤트가 피크 시간대에 몰릴 때 producer가 이를 Kinesis에 PutRecords로 전송하고, CloudWatch는 IncomingRecords, IncomingBytes, PutRecords latency, throttled records 등을 자동으로 수집한다."
    )
    if publish:
        add_table(
            doc,
            ["Metric", "Value"],
            [
                ["Published records", f"{publish['total']:,}"],
                ["Elapsed time", f"{publish['elapsed']:.3f} seconds"],
                ["Overall average publish rate", f"{publish['avg_rps']:.2f} records/sec"],
                ["Failed records after retry", f"{publish['failed_sum']:,}"],
                ["Batches", f"{publish['batches']:,}"],
            ],
            [2.4, 3.6],
        )
    doc.add_paragraph(
        "Producer는 단순히 일정한 속도로 데이터를 밀어 넣지 않았다. 앞의 1.1절에서 정의한 peak-shaped generator와 `--mode peak-shaped`, `--simulated-hour-seconds 30` 설정을 결합해 이벤트 시간대별 볼륨에 따라 재생 속도를 다르게 만들었다. 이 설계 덕분에 CloudWatch 그래프에서 피크 구간의 IncomingRecords가 실제로 튀는 형태로 나타난다. 즉, 테스트 데이터의 event_time만 피크인 것이 아니라 Kinesis ingestion 자체도 피크 형태를 갖는다."
    )
    doc.add_paragraph(
        "단, 이 실험을 해석할 때 가장 중요한 점은 총 281,660건 자체를 production-scale traffic이라고 주장하지 않는 것이다. 실제 대형 서비스에서 20만 건 수준은 충분히 흡수 가능한 규모일 수 있다. 본 실험의 목적은 작은 데이터셋으로 현실 피크를 축소 재현하는 것이며, 이를 위해 event-time의 1시간을 wall-clock 30초로 압축했다. 따라서 평가 기준은 총 record 수가 아니라 records/sec, MiB/sec, throttling 여부, p95/p99 latency, peak/off-peak latency 차이여야 한다."
    )
    add_table(
        doc,
        ["Traffic segment", "Event hours", "Order volume design", "Observed replay implication"],
        [
            ["Off-peak", "00-07, 22-23", "1,666 records/hour", "낮은 ingestion rate"],
            ["Shoulder", "08-10, 14-16, 21", "5,000 records/hour", "중간 수준 ingestion rate"],
            ["Pre/Post peak", "11, 17, 20", "10,000 records/hour", "피크 전후 상승/하강"],
            ["Peak", "12, 13, 18, 19", "50,000 records/hour", "약 893-1,129 records/sec 수준의 높은 replay rate"],
        ],
        [1.2, 1.45, 1.7, 2.1],
        caption="시간대별 주문량 설계와 Kinesis replay 부하 해석",
    )
    add_table(
        doc,
        ["Scale question", "Correct interpretation for this project"],
        [
            ["Is 281,660 records production-scale?", "No. It is a reduced dataset for lab and portfolio validation."],
            ["Why is the experiment still meaningful?", "The replay compresses event time into wall-clock time, so burst pressure appears in throughput metrics."],
            ["What should be evaluated?", "Records/sec, MiB/sec, throttling, failed records after retry, p95/p99 latency, alarm transitions."],
            ["Can the measured 2s p95 be promised at real production scale?", "No. It is a baseline for the tested replay, not a universal SLA."],
            ["How should production capacity be estimated?", "Scale Kinesis shards/on-demand mode, partition key strategy, consumer parallelism, and EMR resources from target peak throughput."],
        ],
        [2.0, 4.2],
        caption="축소 실험 결과를 production claim으로 오해하지 않기 위한 해석 기준",
    )
    doc.add_paragraph(
        "아래 그래프는 교수자가 데이터 분포를 직관적으로 이해할 수 있도록 정제한 주문량 프로파일이다. 낮은 새벽/야간 구간, 완만히 증가하는 오전/오후 구간, 12-13시 점심 피크와 18-19시 저녁 피크가 명확히 구분된다."
    )
    add_fit_picture(doc, chart_path, max_width=6.2, max_height=3.3)
    add_caption(doc, "Figure 2. Hourly order profile used to simulate lunch and dinner peak traffic")
    doc.add_paragraph(
        "Kinesis에서 관찰된 WriteProvisionedThroughputExceeded 또는 throttled record metric은 실패가 아니라 피크 부하를 실제로 밀어 넣었을 때 나타나는 중요한 관찰 지표다. 최종 producer는 partial PutRecords failure를 재시도하도록 수정되었고, 최종 실행에서는 failed_records 합계가 0으로 종료되었다. 따라서 throttling은 시스템이 피크 압력을 받았다는 evidence이고, retry/backoff는 이를 안정적으로 흡수하기 위한 producer-side resilience이다."
    )

    doc.add_heading("7. S3 Raw Zone, Lakehouse Zone, Paimon Table Format", level=1)
    doc.add_paragraph(
        "S3는 이 프로젝트에서 두 가지 zone으로 나뉜다. Raw bucket은 원천 이벤트의 보존 영역이고, Lakehouse bucket은 Paimon warehouse, serving exports, job assets, runtime jars, checkpoint를 보관하는 처리 결과 영역이다. 이 분리는 데이터 엔지니어링에서 매우 중요하다. Raw zone은 immutable하게 유지하고, 처리 로직이 바뀌면 lakehouse/export zone을 재생성할 수 있어야 하기 때문이다."
    )
    add_table(
        doc,
        ["S3 Zone", "Path / Folder", "Purpose"],
        [
            ["Raw bucket", "orders/dt=2026-06-20/order_events.jsonl", "원천 주문 이벤트 JSONL 보관"],
            ["Lakehouse bucket", "paimon/", "Paimon table metadata, snapshot, manifest, data files"],
            ["Lakehouse bucket", "exports/json/", "Frontend dashboard가 읽기 쉬운 serving data"],
            ["Lakehouse bucket", "exports/parquet/", "Glue/Athena-style 분석을 위한 columnar serving data"],
            ["Lakehouse bucket", "jars/", "EMR Serverless에서 사용하는 Paimon Spark runtime jar"],
            ["Lakehouse bucket", "jobs/", "Spark job script assets"],
        ],
        [1.4, 2.4, 2.4],
    )
    doc.add_paragraph(
        "Paimon은 S3 위에서 동작하는 table format으로 이해하면 된다. 즉, Paimon 자체가 별도 database server처럼 떠 있는 것이 아니라 S3에 manifest, snapshot, data file을 저장하고 Spark가 이를 table처럼 읽고 쓰는 구조이다. 이 프로젝트에서는 Paimon을 '최신 상태를 빠르게 반영해야 하는 운영성 current table'을 위한 lakehouse state layer로 사용했다."
    )

    doc.add_heading("8. EMR Serverless Spark 처리 계층", level=1)
    doc.add_paragraph(
        "EMR Serverless는 Spark cluster를 직접 띄우고 관리하지 않고도 Spark job을 실행할 수 있게 해주는 AWS managed service이다. 이 프로젝트에서는 raw validation, Paimon bootstrap, event loading, pressure detection, dashboard export materialization을 EMR Serverless job으로 나누어 실행했다."
    )
    add_table(
        doc,
        ["Job stage", "Input", "Output", "Cloud computing meaning"],
        [
            ["Validate raw events", "S3 raw JSONL", "record count, invalid ratio", "데이터 품질을 cloud job으로 검증"],
            ["Bootstrap Paimon tables", "Paimon jar and Spark catalog config", "table definitions", "외부 table format을 EMR runtime에 통합"],
            ["Load order events", "raw order events", "orders_latest, order_items_latest", "raw event를 lakehouse state로 변환"],
            ["Detect peak pressure", "order/store/hour aggregates", "peak_order_alerts", "운영 알림을 batch/stream hybrid 방식으로 생성"],
            ["Materialize views", "processed tables", "JSON/Parquet exports", "dashboard serving과 analytics table을 분리 제공"],
        ],
        [1.5, 1.4, 1.6, 2.1],
    )
    doc.add_paragraph(
        "EMR Serverless가 중요한 이유는 두 가지다. 첫째, Spark processing capacity를 애플리케이션 코드와 분리할 수 있다. 둘째, VPC/private subnet/IAM/CloudWatch Logs와 결합해 실제 AWS 운영 환경의 조건을 반영할 수 있다. 로컬 pandas script로 결과를 만들 수도 있었지만, 그것은 cloud computing 수업에서 요구하는 managed compute, network isolation, service permission, observability 학습을 충분히 보여주기 어렵다."
    )

    doc.add_heading("9. Glue Data Catalog와 Serving Table 설계", level=1)
    doc.add_paragraph(
        "Glue Data Catalog는 데이터를 저장하는 곳이 아니라, S3에 저장된 데이터의 table 이름, schema, classification, location을 관리하는 metadata catalog이다. 이 프로젝트에서 Glue는 raw JSONL과 Parquet serving exports를 queryable table로 정리하는 역할을 한다."
    )
    add_table(
        doc,
        ["Glue Table", "Format", "S3 Source", "Business/Technical Role"],
        [
            ["raw_order_events_json", "JSON", "raw/orders/", "원천 이벤트 schema 및 raw validation 기준"],
            ["product_demand_daily", "Parquet", "exports/parquet/product_demand_daily/", "PM 상품 수요와 매출 판단"],
            ["product_leaderboard", "Parquet", "exports/parquet/product_leaderboard/", "dashboard 상위 상품 목록"],
            ["store_order_pressure_hourly", "Parquet", "exports/parquet/store_order_pressure_hourly/", "매장/시간대 pressure ratio"],
            ["peak_order_alerts", "Parquet", "exports/parquet/peak_order_alerts/", "PM/SRE alert queue"],
        ],
        [1.45, 0.85, 2.3, 2.0],
    )
    doc.add_paragraph(
        "이 설계는 raw table과 serving table을 분리한다. raw table은 이벤트 원형을 보존하고, serving table은 dashboard와 분석에 필요한 형태로 재구성된다. 따라서 PM dashboard가 빠르게 읽는 JSON export와, Glue catalog가 관리하는 Parquet export를 동시에 제공할 수 있다."
    )

    doc.add_heading("10. VPC, Private Subnet, Endpoint, IAM 권한 설계", level=1)
    doc.add_paragraph(
        "Cloud Computing 프로젝트에서 네트워크와 권한 설정은 부가 요소가 아니라 아키텍처의 신뢰성을 결정하는 핵심이다. 이 프로젝트는 EMR Serverless job이 private subnet 안에서 실행되는 구조를 선택했고, S3와 CloudWatch Logs 접근을 위해 VPC endpoint를 구성했다."
    )
    add_table(
        doc,
        ["Component", "Configuration", "Reason"],
        [
            ["VPC", "peakorder-insight-dev-vpc / 10.42.0.0/16", "프로젝트 전용 네트워크 격리"],
            ["Private subnets", "ap-northeast-2a, ap-northeast-2c", "serverless Spark job의 private execution environment"],
            ["S3 Gateway Endpoint", "com.amazonaws.ap-northeast-2.s3", "private route로 raw/lakehouse bucket 접근"],
            ["CloudWatch Logs Interface Endpoint", "com.amazonaws.ap-northeast-2.logs", "private subnet에서 job log 전송"],
            ["IAM pipeline role", "S3, Glue, Logs, EMR permissions", "EMR job이 필요한 AWS API만 호출하도록 제어"],
        ],
        [1.45, 2.45, 2.25],
    )
    doc.add_paragraph(
        "IAM 설계에서는 S3 read/write, Glue table access, CloudWatch Logs describe/write 권한이 필요했다. 특히 CloudWatch Logs는 단순히 PutLogEvents만 필요한 것이 아니라 log group/stream discovery를 위해 DescribeLogGroups, DescribeLogStreams 권한이 필요할 수 있다. 이 권한이 부족하면 job은 실행되어도 로그를 제대로 확인하기 어렵고, troubleshooting 속도가 급격히 떨어진다."
    )
    add_callout(
        doc,
        "Security interpretation",
        "이 프로젝트는 public network에 모든 것을 열어두는 구조가 아니라, private subnet + VPC endpoint + IAM role 조합으로 AWS managed service에 접근한다. 이는 cloud security와 operational observability를 함께 고려한 설계이다.",
    )

    doc.add_heading("11. CloudWatch 관찰성과 알람 설계", level=1)
    doc.add_paragraph(
        "CloudWatch는 이 프로젝트에서 세 가지 역할을 한다. 첫째, Kinesis throughput과 latency metric을 관찰한다. 둘째, EMR Serverless job log를 수집한다. 셋째, peak ingestion alarm을 통해 피크 유입을 운영 이벤트로 바꾼다."
    )
    add_table(
        doc,
        ["CloudWatch element", "Observed value / setting", "Meaning"],
        [
            ["Kinesis IncomingRecords", "peak graph observed", "피크 replay가 실제 ingestion으로 반영됨"],
            ["PutRecords latency", "write latency visible", "producer write path 성능 관찰"],
            ["Write throughput exceeded", "spike during peak", "provisioned stream에 부하가 걸렸음을 나타냄"],
            ["Alarm threshold", "IncomingRecords >= 40,000 / 1 minute", "피크 구간 자동 감지 기준"],
            ["Log groups", "/aws/emr-serverless, /aws/peakorder/.../pipeline", "job 실행과 troubleshooting evidence"],
        ],
        [1.6, 2.0, 2.6],
    )
    doc.add_paragraph(
        "Alarm의 현재 상태가 OK인 것은 문제가 아니다. 피크가 지나간 뒤 OK로 돌아오는 것이 정상이며, 보고서에서는 threshold line을 넘었던 datapoint와 alarm configuration을 함께 보여주는 것이 중요하다. 이는 피크 구간이 관찰 가능한 운영 이벤트로 모델링되었음을 의미한다."
    )

    doc.add_heading("12. Frontend Dashboard와 PM/SRE 의사결정 지표", level=1)
    doc.add_paragraph(
        "Frontend dashboard는 단순한 시각화 화면이 아니라 lakehouse 결과를 역할별 의사결정으로 바꾸는 serving layer이다. PM, Data Engineer, ML Engineer, SRE, Leadership 탭은 같은 데이터를 서로 다른 질문으로 해석한다."
    )
    add_table(
        doc,
        ["Role", "Primary question", "Representative metric"],
        [
            ["PM", "어떤 피크 구간에 상품/재고/운영 결정을 내려야 하는가?", "peak demand concentration, peak revenue at stake"],
            ["Data Engineer", "파이프라인이 raw부터 export까지 정상적으로 흘렀는가?", "raw count, export tables, freshness"],
            ["ML Engineer", "모델/feature 관점에서 피크 anomaly를 어떻게 활용할 것인가?", "demand drivers, anomaly candidates"],
            ["SRE", "어디서 부하와 지연이 발생하고 알림이 필요한가?", "pressure ratio, CloudWatch alarm, latency"],
            ["Leadership", "피크 대응이 비즈니스 영향과 어떤 관계가 있는가?", "gross sales, alert burden, service risk"],
        ],
        [1.15, 3.1, 2.0],
    )
    doc.add_paragraph(
        "PM KPI decision frame은 기술 metric을 그대로 보여주지 않는다. 예를 들어 Kinesis latency는 PM에게 `Decision freshness`로 제공되고, store-hour aggregate는 `Store pressure ceiling`으로 제공된다. 이는 데이터 엔지니어링 결과가 실제 제품/운영 의사결정 언어로 변환되었음을 보여준다."
    )

    doc.add_heading("13. Latency 분석 결과와 구조 타당성 평가", level=1)
    doc.add_paragraph(
        "Latency 분석은 이 프로젝트의 실시간 구조가 합당했는지 평가하는 핵심 근거이다. 실시간 아키텍처를 사용했다고 주장하려면, 피크 상황에서도 데이터가 충분히 빠르게 관찰되었음을 보여주어야 한다."
    )
    if latency:
        rows = []
        for key in ["all", "peak", "shoulder", "off"]:
            s = latency[key]
            rows.append([
                s["label"],
                f"{int(s['records']):,}",
                f"{s['mean']:.2f} ms",
                f"{s['median']:.2f} ms",
                f"{s['std']:.2f} ms",
                f"{s['p95']:.2f} ms",
                f"{s['p99']:.2f} ms",
                f"{s['max']:.2f} ms",
            ])
        add_table(
            doc,
            ["Segment", "Records", "Mean", "Median", "Std dev", "p95", "p99", "Max"],
            rows,
            [1.0, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8],
        )
    doc.add_paragraph(
        "전체 평균 publish-to-read latency는 약 1.26초, 전체 p95는 약 2.03초였다. 피크 시간대만 보아도 p95는 약 2.06초로 전체와 크게 다르지 않다. 이는 현재 replay 조건에서는 피크 구간에서도 운영 dashboard가 수 초 내 freshness를 유지했다는 뜻이다. 그러나 이 수치는 production traffic 전체에 대한 보장값이 아니라, 이번 압축 replay와 현재 Kinesis capacity 설정에서 얻은 baseline evidence로 해석해야 한다."
    )
    doc.add_paragraph(
        "PM dashboard의 decision freshness 기준을 3초 미만으로 두면, 이번 실험 결과는 기준을 만족한다. 다만 이 결론은 '현실 대규모에서도 무조건 3초 미만'이라는 의미가 아니다. 더 큰 주문량에서는 queueing delay, shard throughput limit, partition key skew, consumer 병렬성 부족, EMR processing delay가 함께 커질 수 있다. 따라서 본 실험은 production SLA 증명이 아니라, 어떤 지표를 관찰하고 어떤 scaling lever를 조정해야 하는지 보여주는 scale-aware validation으로 보는 것이 정확하다."
    )
    add_table(
        doc,
        ["Metric", "Why it matters for scale-up decisions"],
        [
            ["Records/sec", "총량보다 순간 ingestion pressure를 더 잘 표현한다."],
            ["MiB/sec", "Kinesis shard throughput limit과 직접 연결된다."],
            ["WriteProvisionedThroughputExceeded", "현재 shard capacity가 피크 write load를 충분히 받는지 보여준다."],
            ["PutRecords failed/throttled records", "producer retry/backoff와 shard scaling 필요성을 판단한다."],
            ["p95/p99 publish-to-read latency", "운영 dashboard freshness가 tail latency에서도 유지되는지 확인한다."],
            ["Peak/off-peak latency ratio", "피크 상황에서 시스템이 평상시 대비 얼마나 흔들리는지 비교한다."],
        ],
        [2.0, 4.2],
        caption="실제 운영 규모로 확장할 때 latency보다 함께 봐야 하는 capacity metric",
    )
    add_callout(
        doc,
        "Scale-aware evaluation result",
        "Latency evidence는 이 프로젝트의 강한 기술적 근거이지만, production full-scale 보장값으로 과장해서는 안 된다. 더 정확한 해석은 '시간 압축 replay 조건에서 피크형 부하를 재현했고, 현재 capacity에서는 peak-window 200,000건의 p95 latency가 약 2.06초로 관찰되어 dashboard freshness baseline을 확보했다'는 것이다.",
        fill="E7F7EF",
    )

    doc.add_heading("14. 트러블슈팅과 개선 과정", level=1)
    doc.add_paragraph(
        "보고서에서 실패 job은 숨길 대상이 아니라 엔지니어링 사고 과정을 보여주는 증거로 사용하는 것이 좋다. 실제 cloud data platform은 한 번에 완벽히 실행되기보다, network, IAM, runtime dependency, SQL dialect 문제를 하나씩 해결하면서 안정화된다."
    )
    add_table(
        doc,
        ["Issue", "Observed symptom", "Root cause", "Resolution"],
        [
            ["Kinesis partial PutRecords failure", "PutRecords failed for 일부 records", "batch 일부 실패에 대한 retry 부재", "failed records만 exponential backoff로 재시도"],
            ["EMR/Paimon syntax error", "PARSE_SYNTAX_ERROR near '('", "EMR Spark parser와 초기 Paimon DDL syntax 호환성 문제", "append-compatible table definition으로 조정"],
            ["Paimon runtime dependency", "package/runtime resolution issue", "private/serverless 환경에서 external dependency resolve 제약", "Paimon Spark jar를 S3 jars/에 업로드 후 --jars로 주입"],
            ["CloudWatch Logs access", "job log 확인 어려움 또는 AccessDenied 가능성", "private subnet endpoint/IAM describe 권한 필요", "Logs interface endpoint와 logs describe 권한 추가"],
            ["Spark resource tuning", "capacity/resource pressure", "default executor 설정이 실습 환경에 과함", "small executor configuration으로 조정"],
        ],
        [1.3, 1.45, 2.0, 2.0],
    )
    doc.add_paragraph(
        "EMR Studio에 남은 `PARSE_SYNTAX_ERROR`는 보고서에 포함해도 좋다. 단, 이를 단순 문법 실수처럼 쓰지 말고 managed Spark runtime에서 lakehouse table format을 통합할 때 발생한 SQL dialect/runtime compatibility issue로 설명하는 것이 적절하다. 이후 S3 jar 방식과 table definition 조정을 통해 최종 pipeline을 성공시켰다는 흐름이 중요하다."
    )

    doc.add_heading("15. 한계, 비용, 운영 개선안", level=1)
    doc.add_paragraph(
        "현재 프로젝트는 AWS 기반 end-to-end evidence를 확보한 실습/포트폴리오 수준의 구현이다. production 수준으로 확장하려면 다음 개선이 필요하다."
    )
    add_table(
        doc,
        ["Area", "Current implementation", "Production improvement"],
        [
            ["Scale validity", "281,660-record compressed replay", "target production peak RPS/MBps 기준으로 load test 재산정"],
            ["Kinesis scaling", "Provisioned stream with observed throttling", "shard 수 조정, on-demand mode 검토, partition key skew 분석"],
            ["S3 landing", "producer dataset 기반 raw landing", "Kinesis Data Firehose 또는 Lambda consumer로 automatic raw landing"],
            ["Consumer architecture", "local latency observer", "ECS/Lambda/Flink consumer로 지속 실행"],
            ["Data correctness", "append-compatible demo path", "idempotency, deduplication, late event handling"],
            ["Security", "project role and endpoints", "least privilege 세분화, encryption policy, audit logging"],
            ["Cost", "temporary lab resources", "log retention, lifecycle policy, EMR job duration, Kinesis capacity 최적화"],
            ["Dashboard serving", "static/local frontend", "S3 static website, CloudFront, API Gateway, or Amplify hosting"],
        ],
        [1.35, 2.05, 2.7],
    )
    doc.add_paragraph(
        "Load balancing 측면에서는 향후 두 방향으로 확장할 수 있다. 첫째, ingestion load balancing은 Kinesis shard scaling과 partition key 설계를 통해 개선한다. store_id만 partition key로 쓰면 특정 인기 매장에 skew가 발생할 수 있으므로 order_id hash 또는 store_id+time bucket 조합을 고려할 수 있다. 둘째, frontend/API load balancing은 ALB, API Gateway, CloudFront를 사용해 dashboard serving traffic을 분산할 수 있다. 현재 프로젝트는 데이터 ingestion과 processing을 중심으로 구현했기 때문에, ALB 기반 웹 tier는 향후 확장 항목으로 두는 것이 정직하다."
    )

    doc.add_heading("16. 결론", level=1)
    doc.add_paragraph(
        "PeakOrder Insight는 피크 주문이라는 현실적인 운영 문제를 AWS Cloud Computing 관점에서 모델링하고, Kinesis, S3, EMR Serverless, Paimon, Glue, CloudWatch, VPC endpoint, IAM, Terraform을 연결해 구현한 end-to-end 데이터 플랫폼 프로젝트이다."
    )
    doc.add_paragraph(
        "가장 중요한 성과는 세 가지이다. 첫째, 총 281,660건의 주문 이벤트 중 200,000건을 피크 윈도우에 집중시키고 event-time을 wall-clock time보다 빠르게 압축 replay하여 bursty한 ingestion pattern을 만들었다. 둘째, Kinesis producer와 consumer latency evidence를 통해 현재 replay 조건에서 피크 시간대 p95 약 2.06초의 관찰 지연을 확인했다. 셋째, raw data, lakehouse export, Glue schema, CloudWatch alarm, VPC endpoint, dashboard까지 AWS UI와 로그로 검증 가능한 산출물을 확보했다."
    )
    doc.add_paragraph(
        "따라서 이 프로젝트는 단순 AWS 서비스 나열이나 정적 웹사이트 제작이 아니라, cloud-native streaming ingestion, private serverless processing, lakehouse table management, observability, role-based dashboard를 하나의 논리로 묶은 AWS cloud computing case study로 평가할 수 있다. 또한 실험 규모의 한계를 숨기지 않고, records/sec, MiB/sec, shard scaling, partition key, consumer parallelism을 기준으로 production 확장 방향을 제시했다는 점에서 더 현실적인 engineering report로 해석할 수 있다."
    )

    doc.add_page_break()
    doc.add_heading("Appendix A. AWS Console Evidence", level=1)
    doc.add_paragraph(
        "아래 이미지는 구현 결과를 AWS Console 및 local dashboard에서 캡처한 evidence이다. 보고서 제출 형식에 따라 필요한 이미지만 선별하거나, 전체 appendix로 유지할 수 있다."
    )
    for idx, (caption, path_str, description) in enumerate(SCREENSHOTS, 1):
        path = Path(path_str)
        doc.add_heading(caption, level=2)
        if path.exists():
            try:
                add_fit_picture(doc, path, max_height=6.0)
                add_caption(doc, caption)
                add_figure_note(doc, description)
            except Exception as exc:
                doc.add_paragraph(f"Image could not be embedded: {path} ({exc})")
        else:
            doc.add_paragraph(f"Screenshot placeholder: {path}")
            add_figure_note(doc, description)
        if idx not in {len(SCREENSHOTS)}:
            doc.add_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
